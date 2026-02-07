"""
DOCX document processor.
Extracts text content from Microsoft Word documents.
"""
import io
from typing import Optional
from docx import Document
from .base import BaseProcessor


class DOCXProcessor(BaseProcessor):
    """Processor for DOCX documents"""
    
    def extract_text(self, file_content: bytes, filename: str) -> Optional[str]:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content as bytes
            filename: Original filename
            
        Returns:
            Extracted text content, or None if extraction fails
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                return None
            
            full_text = "\n\n".join(text_parts)
            return self.clean_text(full_text)
        
        except Exception as e:
            print(f"Error processing DOCX {filename}: {e}")
            return None
